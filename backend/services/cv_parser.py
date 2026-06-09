from __future__ import annotations

import asyncio
import io
from pathlib import Path

from docx import Document
from docx.table import Table, _Cell
from pypdf import PdfReader


def _extract_pdf_sync(pdf_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


async def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _extract_pdf_sync, pdf_bytes)


def _iter_table_text(table: Table) -> list[str]:
    lines: list[str] = []
    for row in table.rows:
        for cell in row.cells:
            lines.extend(_iter_cell_text(cell))
    return lines


def _iter_cell_text(cell: _Cell) -> list[str]:
    lines = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    for table in cell.tables:
        lines.extend(_iter_table_text(table))
    return lines


def _extract_docx_sync(docx_bytes: bytes) -> str:
    document = Document(io.BytesIO(docx_bytes))
    lines: list[str] = []
    lines.extend(p.text.strip() for p in document.paragraphs if p.text.strip())
    for table in document.tables:
        lines.extend(_iter_table_text(table))
    for section in document.sections:
        for container in (section.header, section.footer):
            lines.extend(p.text.strip() for p in container.paragraphs if p.text.strip())
            for table in container.tables:
                lines.extend(_iter_table_text(table))
    return "\n".join(dict.fromkeys(line for line in lines if line.strip()))


async def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _extract_docx_sync, docx_bytes)


async def extract_text_from_file(path: str) -> str:
    p = Path(path)
    if not p.exists():
        return ""
    if p.suffix.lower() == ".docx":
        return await extract_text_from_docx_bytes(p.read_bytes())
    return await extract_text_from_pdf_bytes(p.read_bytes())
