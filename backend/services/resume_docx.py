from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

from backend.schemas import ResumeTailorerOutput

_BULLET_STYLE = "List Bullet"


def render_resume_docx(output: ResumeTailorerOutput) -> bytes:
    """Render structured resume content to a plain DOCX.

    User text is written as DOCX data via python-docx. This regular-user path
    does not compile templates or edit uploaded files in place.
    """
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    if output.headline:
        heading = document.add_heading(output.headline, level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if output.summary:
        document.add_heading("Summary", level=1)
        document.add_paragraph(output.summary)

    if output.skills:
        document.add_heading("Skills", level=1)
        document.add_paragraph(", ".join(output.skills))

    if output.experience:
        document.add_heading("Experience", level=1)
        for exp in output.experience:
            heading = " - ".join(part for part in [exp.role, exp.company] if part)
            if exp.dates:
                heading = f"{heading} ({exp.dates})" if heading else exp.dates
            if heading:
                document.add_heading(heading, level=2)
            for bullet in exp.bullets:
                if bullet.strip():
                    document.add_paragraph(bullet, style=_BULLET_STYLE)

    if output.projects:
        document.add_heading("Projects", level=1)
        for project in output.projects:
            if project.name:
                document.add_heading(project.name, level=2)
            if project.description:
                document.add_paragraph(project.description)
            for bullet in project.bullets:
                if bullet.strip():
                    document.add_paragraph(bullet, style=_BULLET_STYLE)

    if output.education:
        document.add_heading("Education", level=1)
        for education in output.education:
            line = " - ".join(part for part in [education.degree, education.institution] if part)
            if education.dates:
                line = f"{line} ({education.dates})" if line else education.dates
            if line:
                document.add_paragraph(line)

    # NOTE: output.tailored_bullets is deliberately NOT rendered. It is metadata
    # explaining which bullets were rewritten; the rewrites are already placed in
    # the Experience/Projects sections above, so rendering it again ("Additional
    # Highlights") duplicated content verbatim and pushed the document to 2 pages.

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
