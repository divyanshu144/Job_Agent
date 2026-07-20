from unittest.mock import MagicMock, patch

from docx import Document


async def test_extract_text_from_bytes():
    mock_reader = MagicMock()
    mock_reader.pages = [
        MagicMock(extract_text=MagicMock(return_value="Software Engineer")),
        MagicMock(extract_text=MagicMock(return_value="Python FastAPI")),
    ]
    with patch("backend.services.cv_parser.PdfReader", return_value=mock_reader):
        from backend.services.cv_parser import extract_text_from_pdf_bytes

        result = await extract_text_from_pdf_bytes(b"fake-pdf-bytes")
    assert "Software Engineer" in result
    assert "Python FastAPI" in result


async def test_extract_text_falls_back_to_ocr_for_scanned_pdf():
    mock_reader = MagicMock()
    mock_reader.pages = [MagicMock(extract_text=MagicMock(return_value=""))]
    with (
        patch("backend.services.cv_parser.PdfReader", return_value=mock_reader),
        patch("backend.services.cv_parser.convert_from_bytes", return_value=["page-1-image"]),
        patch(
            "backend.services.cv_parser.pytesseract.image_to_string",
            return_value="OCR-recovered resume text with Python experience.",
        ),
    ):
        from backend.services.cv_parser import extract_text_from_pdf_bytes

        result = await extract_text_from_pdf_bytes(b"scanned-pdf-bytes")
    assert "OCR-recovered resume text" in result


async def test_extract_text_keeps_pypdf_result_when_ocr_fails():
    mock_reader = MagicMock()
    mock_reader.pages = [MagicMock(extract_text=MagicMock(return_value=""))]
    with (
        patch("backend.services.cv_parser.PdfReader", return_value=mock_reader),
        patch(
            "backend.services.cv_parser.convert_from_bytes",
            side_effect=RuntimeError("poppler not installed"),
        ),
    ):
        from backend.services.cv_parser import extract_text_from_pdf_bytes

        result = await extract_text_from_pdf_bytes(b"scanned-pdf-bytes")
    assert result == ""


async def test_extract_text_skips_ocr_when_pypdf_already_has_enough_text():
    mock_reader = MagicMock()
    mock_reader.pages = [
        MagicMock(extract_text=MagicMock(return_value="Plenty of real extracted text here.")),
    ]
    with (
        patch("backend.services.cv_parser.PdfReader", return_value=mock_reader),
        patch("backend.services.cv_parser.convert_from_bytes") as mock_convert,
    ):
        from backend.services.cv_parser import extract_text_from_pdf_bytes

        result = await extract_text_from_pdf_bytes(b"real-pdf-bytes")
    assert result == "Plenty of real extracted text here."
    mock_convert.assert_not_called()


async def test_extract_text_missing_file_returns_empty():
    from backend.services.cv_parser import extract_text_from_file

    result = await extract_text_from_file("nonexistent/path.pdf")
    assert result == ""


async def test_extract_docx_walks_tables_and_headers(tmp_path):
    from backend.services.cv_parser import extract_text_from_docx_bytes

    path = tmp_path / "resume.docx"
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Header Candidate"
    document.add_paragraph("Top-level Summary")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "Postgres"
    document.save(path)

    result = await extract_text_from_docx_bytes(path.read_bytes())

    assert "Header Candidate" in result
    assert "Top-level Summary" in result
    assert "Python" in result
    assert "Postgres" in result
